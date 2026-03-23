#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档格式修正 API 服务
使用 Flask 提供 HTTP API 接口
"""

from flask import Flask, request, jsonify, send_file
from docx import Document
import os
import tempfile

app = Flask(__name__)


@app.route('/')
def index():
    """API 首页"""
    return jsonify({
        'service': 'Word格式修正API',
        'version': '1.0',
        'endpoint': '/format',
        'method': 'POST',
        'parameters': ['template', 'input']
    })


@app.route('/format', methods=['POST'])
def format_document():
    """处理文档格式修正请求"""
    try:
        # 尝试从 JSON 获取 base64 编码的文件
        if request.is_json:
            data = request.get_json()
            template_b64 = data.get('template')
            input_b64 = data.get('input')

            if not template_b64 or not input_b64:
                return jsonify({'error': '缺少 template 或 input 参数'}), 400

            # 解码 base64
            import base64
            template_data = base64.b64decode(template_b64)
            input_data = base64.b64decode(input_b64)

            # 保存临时文件
            tmp_dir = tempfile.gettempdir()
            template_path = os.path.join(tmp_dir, 'template.docx')
            input_path = os.path.join(tmp_dir, 'input.docx')
            output_path = os.path.join(tmp_dir, 'output.docx')

            with open(template_path, 'wb') as f:
                f.write(template_data)
            with open(input_path, 'wb') as f:
                f.write(input_data)

        else:
            # 兼容文件上传方式
            template_file = request.files.get('template')
            input_file = request.files.get('input')

            if not template_file or not input_file:
                return jsonify({'error': '需要上传两个文件或提供 base64 数据'}), 400

            # 保存临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_template:
                template_file.save(tmp_template.name)
                template_path = tmp_template.name

            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_input:
                input_file.save(tmp_input.name)
                input_path = tmp_input.name

            output_path = tempfile.mktemp(suffix='.docx')

        # 提取格式规则
        rules = extract_format_rules(template_path)

        # 应用格式
        apply_format(input_path, output_path, rules)

        # 清理临时文件
        os.unlink(template_path)
        os.unlink(input_path)

        # 如果是 JSON 请求，返回 base64 编码的文件
        if request.is_json:
            import base64
            with open(output_path, 'rb') as f:
                result_data = base64.b64encode(f.read()).decode()
            os.unlink(output_path)
            return jsonify({'file': result_data, 'message': '格式修正完成'})
        else:
            # 返回文件下载
            return send_file(output_path, as_attachment=True, download_name='formatted.docx')

    except Exception as e:
        return jsonify({'error': str(e)}), 500


def extract_format_rules(template_path):
    """提取格式规则"""
    doc = Document(template_path)
    rules = {}
    if doc.paragraphs and doc.paragraphs[0].runs:
        run = doc.paragraphs[0].runs[0]
        para = doc.paragraphs[0]
        rules['font_name'] = run.font.name
        rules['font_size'] = run.font.size
        rules['alignment'] = para.alignment
        rules['line_spacing'] = para.paragraph_format.line_spacing
    return rules


def apply_format(input_path, output_path, rules):
    """应用格式"""
    doc = Document(input_path)
    for para in doc.paragraphs:
        if rules.get('alignment'):
            para.alignment = rules['alignment']
        if rules.get('line_spacing'):
            para.paragraph_format.line_spacing = rules['line_spacing']
        for run in para.runs:
            if rules.get('font_name'):
                run.font.name = rules['font_name']
            if rules.get('font_size'):
                run.font.size = rules['font_size']
    doc.save(output_path)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
